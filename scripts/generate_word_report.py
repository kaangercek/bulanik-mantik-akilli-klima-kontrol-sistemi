from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
REPORT_MD = ROOT / "REPORT.md"
OUTPUT_DOCX = ROOT / "Bulanik_Mantik_Akilli_Klima_Kontrol_Sistemi_Raporu.docx"


TITLE = "Bulanık Mantık Tabanlı Akıllı Klima Kontrol Sistemi"
SUBTITLE = "Bulanık Mantık Dersi Dönem Projesi Raporu"


def set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal_style.font.size = Pt(12)

    for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    document.styles["Title"].font.size = Pt(20)
    document.styles["Title"].font.bold = True
    document.styles["Subtitle"].font.size = Pt(14)
    document.styles["Heading 1"].font.size = Pt(15)
    document.styles["Heading 1"].font.bold = True
    document.styles["Heading 2"].font.size = Pt(13)
    document.styles["Heading 2"].font.bold = True
    document.styles["Heading 3"].font.size = Pt(12)
    document.styles["Heading 3"].font.bold = True


def set_cell_text(cell, text: str, bold: bool = False, center: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "İçindekiler için Word içinde sağ tıklayıp 'Alanı Güncelleştir' seçebilirsin."

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(placeholder)
    run._r.append(fld_char_end)


def add_footer_with_page_number(document: Document) -> None:
    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(paragraph)


def add_cover_page(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.space_after = Pt(18)
    paragraph.space_before = Pt(120)
    run = paragraph.add_run(TITLE)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(20)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.space_after = Pt(80)
    run = subtitle.add_run(SUBTITLE)
    run.italic = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)

    info_rows = [
        ("Ders", "Bulanık Mantık"),
        ("Proje Konusu", "Akıllı Klima / Oda Konfor Kontrol Sistemi"),
        ("Hazırlayan", "Ad Soyad"),
        ("Öğrenci No", "Numaranızı yazınız"),
        ("Teslim Tarihi", "21.05.2026"),
    ]

    table = document.add_table(rows=len(info_rows), cols=2)
    table.style = "Table Grid"
    table.autofit = True
    for row, (label, value) in zip(table.rows, info_rows):
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_text(row.cells[1], value)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.space_before = Pt(70)
    note_run = note.add_run("Not: Kapak sayfasındaki kişisel bilgileri teslimden önce güncelleyiniz.")
    note_run.italic = True
    note_run.font.name = "Times New Roman"
    note_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    note_run.font.size = Pt(11)

    document.add_page_break()


def add_toc_page(document: Document) -> None:
    heading = document.add_paragraph(style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run("İçindekiler")

    toc_paragraph = document.add_paragraph()
    add_toc(toc_paragraph)
    document.add_page_break()


def add_run_with_markdown(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)")
    parts = pattern.split(text)

    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)

        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(11)
        else:
            run.text = part


def add_normal_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    add_run_with_markdown(paragraph, text)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(3)
    add_run_with_markdown(paragraph, text)


def add_numbered(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(3)
    add_run_with_markdown(paragraph, text)


def parse_table(lines: list[str], start_index: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start_index
    while index < len(lines) and lines[index].strip().startswith("|"):
        stripped = lines[index].strip()
        if not re.fullmatch(r"\|[\-\s:|]+\|", stripped):
            cells = [cell.strip() for cell in stripped.strip("|").split("|")]
            rows.append(cells)
        index += 1
    return rows, index


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"

    for row_index, row in enumerate(rows):
        for col_index in range(column_count):
            value = row[col_index] if col_index < len(row) else ""
            set_cell_text(
                table.rows[row_index].cells[col_index],
                value,
                bold=(row_index == 0),
                center=(row_index == 0 or value.replace(".", "", 1).isdigit()),
            )

    document.add_paragraph()


def add_markdown_body(document: Document, markdown_text: str) -> None:
    lines = markdown_text.splitlines()
    index = 0
    skipped_title = False

    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()

        if not line:
            index += 1
            continue

        if line.startswith("# "):
            if not skipped_title:
                skipped_title = True
            else:
                paragraph = document.add_paragraph(style="Heading 1")
                paragraph.paragraph_format.space_before = Pt(12)
                paragraph.paragraph_format.space_after = Pt(6)
                paragraph.add_run(line[2:].strip())
            index += 1
            continue

        if line.startswith("## "):
            paragraph = document.add_paragraph(style="Heading 1")
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.add_run(line[3:].strip())
            index += 1
            continue

        if line.startswith("### "):
            paragraph = document.add_paragraph(style="Heading 2")
            paragraph.paragraph_format.space_before = Pt(10)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.add_run(line[4:].strip())
            index += 1
            continue

        if line.startswith("#### "):
            paragraph = document.add_paragraph(style="Heading 3")
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.add_run(line[5:].strip())
            index += 1
            continue

        if line.startswith("|"):
            table_rows, index = parse_table(lines, index)
            add_table(document, table_rows)
            continue

        if re.match(r"^- ", line):
            add_bullet(document, line[2:].strip())
            index += 1
            continue

        if re.match(r"^\d+\.\s", line):
            numbered_text = re.sub(r"^\d+\.\s*", "", line)
            add_numbered(document, numbered_text)
            index += 1
            continue

        add_normal_paragraph(document, line)
        index += 1


def generate_docx() -> Path:
    markdown_text = REPORT_MD.read_text(encoding="utf-8")
    document = Document()
    set_document_defaults(document)
    add_cover_page(document)
    add_toc_page(document)
    add_markdown_body(document, markdown_text)
    add_footer_with_page_number(document)
    document.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    path = generate_docx()
    print(path)
